from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0x10, 0x6B, 0x3E)  # vert emerald
    h.font.name = "Calibri"

PRIMARY = RGBColor(0x10, 0x6B, 0x3E)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = RGBColor(0x10, 0x6B, 0x3E)
ALT_ROW = RGBColor(0xF0, 0xF9, 0xF4)


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): str(color),
    })
    shading.append(shd)


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MallHub")
run.font.size = Pt(42)
run.font.color.rgb = PRIMARY
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Votre centre commercial en ligne")
run.font.size = Pt(16)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Documentation technique et fonctionnelle")
run.font.size = Pt(14)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Projet Final — MEAN Stack\nMaster 1 — Web Avancé\n2024-2025")
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

member_title = doc.add_paragraph()
member_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = member_title.add_run("Membres du groupe")
run.font.size = Pt(14)
run.bold = True

members = [
    ("1", "RAJOSOA Tsiory Diary Luc", "ETU002665"),
    ("2", "RAMAROZATOVO Tahiry Kevin", "ETU002546"),
]
add_styled_table(["N°", "Nom et Prénom(s)", "N° Étudiant"], members, [2, 8, 4])

repo = doc.add_paragraph()
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = repo.add_run("Dépôt Git : ")
run.font.size = Pt(11)
run = repo.add_run("https://github.com/Kevin-rm/m1p13mean-Diary-Kevin")
run.font.size = Pt(11)
run.font.color.rgb = PRIMARY
run.underline = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table des matières", level=1)
toc_items = [
    "1. Présentation de l'application",
    "   1.1. Vue d'ensemble",
    "   1.2. Fonctionnalités principales",
    "   1.3. Architecture technique",
    "   1.4. Description des pages",
    "2. Structure MongoDB",
    "   2.1. Schéma global",
    "   2.2. Détail des collections",
    "3. Liste des URL de l'API",
    "   3.1. Routes publiques",
    "   3.2. Routes authentifiées",
    "   3.3. Routes par module",
    "4. Informations du projet",
    "   4.1. Lien Git et branches",
    "   4.2. Membres du groupe",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith("   "):
        p.runs[0].bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. PRÉSENTATION DE L'APPLICATION
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Présentation de l'application", level=1)

doc.add_heading("1.1. Vue d'ensemble", level=2)
doc.add_paragraph(
    "MallHub est une plateforme e-commerce de type marketplace (centre commercial en ligne) "
    "développée avec le stack MEAN (MongoDB, Express.js, Angular, Node.js). "
    "Elle permet à plusieurs boutiques de vendre leurs produits sur une plateforme unique, "
    "offrant aux clients une expérience d'achat centralisée."
)
doc.add_paragraph(
    "L'application se compose de trois espaces principaux :"
)

items = [
    ("Frontoffice (Client)", "Interface publique et espace client pour parcourir les boutiques, "
     "consulter les produits, gérer le panier, passer des commandes, gérer les favoris et laisser des avis."),
    ("Backoffice Boutique", "Tableau de bord pour les propriétaires et membres de boutiques : "
     "gestion des produits, promotions, mouvements de stock, commandes et membres de l'équipe."),
    ("Backoffice Admin", "Interface d'administration pour gérer les boutiques (validation, suspension), "
     "les catégories de produits et visualiser les statistiques globales."),
]
for title_text, desc_text in items:
    p = doc.add_paragraph()
    run = p.add_run(f"• {title_text} : ")
    run.bold = True
    p.add_run(desc_text)

doc.add_heading("1.2. Fonctionnalités principales", level=2)

doc.add_heading("Authentification et comptes", level=3)
features_auth = [
    "Inscription client ou propriétaire de boutique",
    "Connexion / déconnexion avec gestion de sessions JWT (access + refresh tokens)",
    "Gestion du profil utilisateur (nom, email, avatar, mot de passe)",
    "Système d'invitations pour rejoindre une boutique (rôles : owner, manager, seller)",
]
for f in features_auth:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Navigation et découverte (Frontoffice)", level=3)
features_front = [
    "Page d'accueil avec carrousel de produits, catégories et boutiques populaires",
    "Catalogue produits avec filtres (catégorie, prix, tri) et recherche",
    "Catalogue boutiques avec tri par note et recherche",
    "Page détail produit avec carrousel d'images, informations complètes et actions",
    "Page détail boutique avec galerie, horaires, informations de contact et liste des produits",
]
for f in features_front:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Panier et commandes", level=3)
features_orders = [
    "Panier persistant (localStorage) avec ajout, modification de quantité et suppression",
    "Checkout multi-boutiques : une commande par boutique avec un checkoutRef partagé",
    "Historique des commandes regroupées par checkout",
    "Suivi de statut : en attente → confirmée / refusée / annulée",
]
for f in features_orders:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Favoris", level=3)
features_fav = [
    "Ajout / suppression de produits en favoris",
    "Liste de favoris avec actions rapides (ajouter au panier, supprimer)",
]
for f in features_fav:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Avis et notes", level=3)
features_reviews = [
    "Système d'avis par boutique (1 avis par utilisateur par boutique)",
    "Note de 1 à 5 étoiles avec commentaire optionnel",
    "Calcul automatique de la note moyenne et du nombre total d'avis",
    "Affichage des avis sur la page boutique avec pagination",
]
for f in features_reviews:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Gestion de boutique (Backoffice)", level=3)
features_shop = [
    "Tableau de bord avec statistiques (commandes, CA, produits, stock faible)",
    "Gestion des produits (CRUD, images, activation/désactivation)",
    "Gestion des promotions (pourcentage ou montant fixe)",
    "Mouvements de stock (entrée, sortie, ajustement)",
    "Gestion des commandes (confirmation, refus avec motif, annulation)",
    "Gestion des membres et invitations (owner, manager, seller)",
    "Paramètres boutique (logo, images, horaires, contact, description)",
]
for f in features_shop:
    doc.add_paragraph(f"• {f}")

doc.add_heading("Administration (Backoffice Admin)", level=3)
features_admin = [
    "Tableau de bord avec statistiques globales (boutiques, commandes, CA)",
    "Validation / suspension des boutiques",
    "Gestion des catégories de produits (CRUD)",
]
for f in features_admin:
    doc.add_paragraph(f"• {f}")

doc.add_heading("1.3. Architecture technique", level=2)
add_styled_table(
    ["Couche", "Technologie", "Description"],
    [
        ["Frontend", "Angular 21", "SPA avec composants standalone, Signals, TanStack Query"],
        ["UI", "PrimeNG 21 + Tailwind CSS 4", "Composants PrimeNG (preset Aura) + utilitaires Tailwind"],
        ["Backend", "Express.js + Node.js", "API REST avec architecture modulaire (routes → controllers → services)"],
        ["Base de données", "MongoDB + Mongoose", "ODM avec schémas, index, TTL, validations"],
        ["Auth", "JWT (cookies HttpOnly)", "Access token + Refresh token, middleware authenticate/authorize"],
        ["Package manager", "pnpm", "Gestionnaire de paquets performant"],
    ],
    [4, 5, 9],
)

doc.add_heading("1.4. Description des pages", level=2)

doc.add_heading("Pages publiques", level=3)
add_styled_table(
    ["Page", "URL", "Description"],
    [
        ["Landing", "/", "Page d'accueil avec hero, carrousel produits, catégories, boutiques populaires"],
        ["Boutiques", "/shops", "Liste des boutiques avec filtres, tri et recherche"],
        ["Détail boutique", "/shops/:id", "Détail complet d'une boutique avec produits et avis"],
        ["Produits", "/products", "Catalogue avec filtres (catégorie, prix, tri) et recherche"],
        ["Détail produit", "/products/:id", "Carrousel d'images, info produit, actions panier/favoris"],
        ["Connexion", "/login", "Formulaire de connexion"],
        ["Inscription", "/register", "Formulaire d'inscription (client ou boutique)"],
    ],
    [4, 4, 10],
)

doc.add_heading("Pages client (authentifié)", level=3)
add_styled_table(
    ["Page", "URL", "Description"],
    [
        ["Accueil", "/home", "Tableau de bord client : raccourcis, commandes récentes, produits"],
        ["Panier", "/cart", "Gestion du panier avec récapitulatif et checkout"],
        ["Favoris", "/favorites", "Liste des produits favoris"],
        ["Commandes", "/orders", "Historique des commandes groupées par checkout"],
        ["Profil", "/profile", "Modification profil, avatar, mot de passe"],
    ],
    [4, 4, 10],
)

doc.add_heading("Backoffice boutique", level=3)
add_styled_table(
    ["Page", "URL", "Description"],
    [
        ["Dashboard", "/backoffice/shop/dashboard", "Stats : commandes, CA, produits, stock faible"],
        ["Ma boutique", "/backoffice/shop/my-shop", "Paramètres boutique (logo, images, horaires)"],
        ["Produits", "/backoffice/shop/products", "Liste et gestion des produits"],
        ["Nouveau produit", "/backoffice/shop/products/new", "Création de produit avec images"],
        ["Mouvements stock", "/backoffice/shop/stock-movements", "Historique des mouvements de stock"],
        ["Commandes", "/backoffice/shop/orders", "Gestion des commandes reçues"],
        ["Membres", "/backoffice/shop/members", "Gestion de l'équipe et invitations"],
    ],
    [4, 6, 8],
)

doc.add_heading("Backoffice admin", level=3)
add_styled_table(
    ["Page", "URL", "Description"],
    [
        ["Dashboard", "/backoffice/admin/dashboard", "Stats globales : boutiques, commandes, CA"],
        ["Boutiques", "/backoffice/admin/shops", "Validation / suspension des boutiques"],
        ["Catégories", "/backoffice/admin/categories", "CRUD des catégories de produits"],
    ],
    [4, 6, 8],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. STRUCTURE MONGODB
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Structure MongoDB", level=1)

doc.add_heading("2.1. Schéma global", level=2)
doc.add_paragraph(
    "La base de données MallHub comporte 15 collections organisées en modules fonctionnels. "
    "Toutes les collections utilisent les timestamps automatiques (createdAt, updatedAt) "
    "et une transformation JSON qui renomme _id en id et supprime __v."
)

add_styled_table(
    ["Module", "Collection", "Description"],
    [
        ["Utilisateurs", "users", "Comptes utilisateurs"],
        ["Utilisateurs", "profiles", "Profils (admin, shop, customer)"],
        ["Utilisateurs", "roles", "Rôles boutique (owner, manager, seller)"],
        ["Utilisateurs", "usercontexts", "Liaison user ↔ profile ↔ role ↔ shop"],
        ["Auth", "refreshtokens", "Tokens de rafraîchissement (TTL)"],
        ["Boutiques", "shops", "Boutiques avec horaires, position, note"],
        ["Boutiques", "invitations", "Invitations à rejoindre une boutique"],
        ["Catalogue", "categories", "Catégories de produits"],
        ["Catalogue", "products", "Produits avec images et stock"],
        ["Catalogue", "promotions", "Promotions (% ou montant fixe)"],
        ["Catalogue", "stockmovements", "Mouvements de stock"],
        ["Commandes", "orders", "Commandes avec items et historique de statut"],
        ["Commandes", "carts", "Paniers utilisateur"],
        ["Social", "reviews", "Avis boutique (1-5 étoiles)"],
        ["Social", "favorites", "Favoris (produits/boutiques)"],
    ],
    [3, 4, 11],
)

doc.add_heading("2.2. Détail des collections", level=2)

# ── Users ──
doc.add_heading("Collection : users", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["firstName", "String", "Requis, trimmed"],
        ["lastName", "String", "Requis, trimmed"],
        ["email", "String", "Requis, unique, lowercase"],
        ["password", "String", "Requis, min 8 chars, non retourné par défaut"],
        ["avatarUrl", "String", "Optionnel"],
        ["isActive", "Boolean", "Défaut : true"],
        ["lastLoginAt", "Date", "Optionnel"],
    ],
    [4, 4, 10],
)

# ── Profiles ──
doc.add_heading("Collection : profiles", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["code", "String", "Requis, unique, enum: admin | shop | customer"],
        ["label", "String", "Requis"],
        ["description", "String", "Optionnel"],
        ["permissions", "[String]", "Liste de permissions (voir ci-dessous)"],
    ],
    [4, 4, 10],
)

p = doc.add_paragraph()
run = p.add_run("Permissions disponibles : ")
run.bold = True
doc.add_paragraph(
    "shops:create, shops:validate, shops:suspend, shops:manage, "
    "categories:read, categories:write, "
    "products:read, products:write, "
    "orders:read, orders:manage, orders:create, "
    "members:read, members:manage, "
    "stats:read, stats:global, "
    "moderation:reviews, moderation:products, "
    "cart:manage, reviews:write, favorites:manage"
)

# ── Roles ──
doc.add_heading("Collection : roles", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["code", "String", "Requis, unique, enum: owner | manager | seller"],
        ["label", "String", "Requis"],
        ["description", "String", "Optionnel"],
        ["permissions", "[String]", "Sous-ensemble des permissions"],
    ],
    [4, 4, 10],
)

# ── UserContexts ──
doc.add_heading("Collection : usercontexts", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["user", "ObjectId → User", "Requis"],
        ["profile", "ObjectId → Profile", "Requis"],
        ["role", "ObjectId → Role", "Optionnel (null pour admin/customer)"],
        ["shop", "ObjectId → Shop", "Optionnel (null pour admin/customer)"],
        ["isActive", "Boolean", "Défaut : true"],
    ],
    [4, 5, 9],
)
doc.add_paragraph("Index unique : { user, profile, shop }")

# ── RefreshTokens ──
doc.add_heading("Collection : refreshtokens", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["token", "String", "Requis, unique"],
        ["user", "ObjectId → User", "Requis"],
        ["userContext", "ObjectId → UserContext", "Requis"],
        ["expiresAt", "Date", "Requis, TTL index"],
        ["revokedAt", "Date", "Optionnel"],
    ],
    [4, 5, 9],
)

# ── Shops ──
doc.add_heading("Collection : shops", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["name", "String", "Requis, unique, trimmed"],
        ["description", "String", "Requis, trimmed"],
        ["logoUrl", "String", "Optionnel"],
        ["images", "[String]", "Défaut : []"],
        ["status", "String", "enum: pending | active | suspended"],
        ["contactEmail", "String", "Optionnel, lowercase"],
        ["contactPhone", "String", "Optionnel"],
        ["schedule", "[{day, openTime, closeTime}]", "Horaires d'ouverture"],
        ["position", "{x, y, floor}", "Position dans le mall"],
        ["averageRating", "Number", "0-5, défaut 0"],
        ["totalReviews", "Number", "Défaut 0"],
        ["owner", "ObjectId → User", "Requis"],
    ],
    [4, 5, 9],
)

# ── Categories ──
doc.add_heading("Collection : categories", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["name", "String", "Requis, unique, trimmed"],
        ["description", "String", "Optionnel"],
        ["imageUrl", "String", "Optionnel"],
        ["isActive", "Boolean", "Défaut : true"],
    ],
    [4, 4, 10],
)

# ── Products ──
doc.add_heading("Collection : products", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["name", "String", "Requis, trimmed"],
        ["description", "String", "Optionnel"],
        ["price", "Number", "Requis, min 0"],
        ["stock", "Number", "Requis, min 0, défaut 0"],
        ["images", "[String]", "Défaut : []"],
        ["isActive", "Boolean", "Défaut : true"],
        ["shop", "ObjectId → Shop", "Requis"],
        ["category", "ObjectId → Category", "Requis"],
    ],
    [4, 5, 9],
)
doc.add_paragraph("Index unique : { shop, name }")

# ── Promotions ──
doc.add_heading("Collection : promotions", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["product", "ObjectId → Product", "Requis"],
        ["shop", "ObjectId → Shop", "Requis"],
        ["type", "String", "enum: percentage | fixed"],
        ["value", "Number", "Requis, min 0"],
        ["startDate", "Date", "Requis"],
        ["endDate", "Date", "Requis"],
        ["isActive", "Boolean", "Défaut : true"],
    ],
    [4, 5, 9],
)

# ── StockMovements ──
doc.add_heading("Collection : stockmovements", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["shop", "ObjectId → Shop", "Requis"],
        ["performedBy", "ObjectId → User", "Requis"],
        ["date", "Date", "Requis, défaut now"],
        ["type", "String", "enum: in | out | adjustment"],
        ["note", "String", "Optionnel"],
        ["lines", "[{product, quantity, previousStock, newStock}]", "Requis, min 1 ligne"],
        ["lineCount", "Number", "Requis, min 1"],
    ],
    [4, 6, 8],
)

# ── Orders ──
doc.add_heading("Collection : orders", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["orderNumber", "String", "Requis, unique"],
        ["buyer", "ObjectId → User", "Requis"],
        ["shop", "ObjectId → Shop", "Requis"],
        ["items", "[{product, productName, productPrice, productImageUrl, quantity, subtotal}]", "Requis, min 1"],
        ["totalAmount", "Number", "Requis, min 0"],
        ["status", "String", "enum: pending | confirmed | refused | cancelled"],
        ["statusHistory", "[{from, to, changedBy, changedAt, reason}]", "Historique des transitions"],
        ["note", "String", "Optionnel"],
        ["checkoutRef", "String", "Groupement de commandes d'un même checkout"],
    ],
    [4, 6, 8],
)

# ── Carts ──
doc.add_heading("Collection : carts", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["user", "ObjectId → User", "Requis, unique"],
        ["items", "[{product, shop, quantity, addedAt}]", "Défaut : []"],
    ],
    [4, 6, 8],
)

# ── Reviews ──
doc.add_heading("Collection : reviews", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["user", "ObjectId → User", "Requis"],
        ["shop", "ObjectId → Shop", "Requis"],
        ["rating", "Number", "Requis, 1-5"],
        ["comment", "String", "Optionnel, max 1000 chars"],
        ["status", "String", "enum: pending | approved | rejected"],
        ["moderatedBy", "ObjectId → User", "Optionnel"],
        ["moderatedAt", "Date", "Optionnel"],
    ],
    [4, 5, 9],
)
doc.add_paragraph("Index unique : { user, shop } — un seul avis par utilisateur par boutique")

# ── Favorites ──
doc.add_heading("Collection : favorites", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["user", "ObjectId → User", "Requis"],
        ["targetType", "String", "enum: Shop | Product"],
        ["target", "ObjectId (polymorphe)", "Réf dynamique via targetType"],
    ],
    [4, 5, 9],
)

# ── Invitations ──
doc.add_heading("Collection : invitations", level=3)
add_styled_table(
    ["Champ", "Type", "Contraintes"],
    [
        ["shop", "ObjectId → Shop", "Requis"],
        ["user", "ObjectId → User", "Requis"],
        ["role", "ObjectId → Role", "Requis"],
        ["status", "String", "enum: pending | accepted | declined | cancelled"],
        ["invitedBy", "ObjectId → User", "Requis"],
    ],
    [4, 5, 9],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. LISTE DES URL DE L'API
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Liste des URL de l'API", level=1)
doc.add_paragraph("Base URL : /api")

doc.add_heading("3.1. Routes publiques (sans authentification)", level=2)

add_styled_table(
    ["Méthode", "URL", "Description"],
    [
        ["GET", "/api/health", "Health check"],
        ["POST", "/api/auth/register/customer", "Inscription client"],
        ["POST", "/api/auth/register/shop", "Inscription propriétaire de boutique"],
        ["POST", "/api/auth/login", "Connexion"],
        ["POST", "/api/auth/logout", "Déconnexion"],
        ["POST", "/api/auth/refresh", "Rafraîchir le token"],
        ["GET", "/api/public/shops", "Liste des boutiques actives"],
        ["GET", "/api/public/shops/:id", "Détail d'une boutique"],
        ["GET", "/api/public/shops/:id/reviews", "Avis d'une boutique"],
        ["GET", "/api/public/shops/:id/products", "Produits d'une boutique"],
        ["GET", "/api/public/products", "Liste des produits actifs"],
        ["GET", "/api/public/products/:id", "Détail d'un produit"],
        ["GET", "/api/public/categories", "Liste des catégories actives"],
    ],
    [2, 7, 9],
)

doc.add_heading("3.2. Routes authentifiées (auth requise)", level=2)

doc.add_heading("Compte utilisateur — /api/account", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["PATCH", "/api/account/profile", "—", "Modifier le profil"],
        ["PATCH", "/api/account/avatar", "—", "Modifier l'avatar"],
        ["PATCH", "/api/account/password", "—", "Changer le mot de passe"],
        ["GET", "/api/account/invitations", "—", "Mes invitations reçues"],
        ["POST", "/api/account/invitations/:id/accept", "—", "Accepter une invitation"],
        ["POST", "/api/account/invitations/:id/decline", "—", "Décliner une invitation"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Utilisateur authentifié — /api/auth", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/auth/me", "—", "Informations utilisateur courant"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("3.3. Routes par module", level=2)

doc.add_heading("Catégories — /api/categories", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/categories", "categories:read", "Liste paginée"],
        ["GET", "/api/categories/select", "—", "Dropdown (select)"],
        ["GET", "/api/categories/:id", "categories:read", "Détail"],
        ["POST", "/api/categories", "categories:write", "Créer"],
        ["PATCH", "/api/categories/:id", "categories:write", "Modifier"],
        ["PATCH", "/api/categories/:id/toggle-active", "categories:write", "Activer/désactiver"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Produits — /api/products", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/products", "products:read", "Liste paginée, filtrable"],
        ["GET", "/api/products/select", "products:read", "Dropdown (select)"],
        ["GET", "/api/products/stats", "products:read", "Statistiques"],
        ["GET", "/api/products/:id", "products:read", "Détail"],
        ["POST", "/api/products", "products:write", "Créer (avec images)"],
        ["PATCH", "/api/products/:id", "products:write", "Modifier"],
        ["PATCH", "/api/products/:id/toggle-active", "products:write", "Activer/désactiver"],
        ["POST", "/api/products/:id/images", "products:write", "Ajouter des images"],
        ["DELETE", "/api/products/:id/images", "products:write", "Supprimer une image"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Promotions — /api/promotions", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/promotions", "products:read", "Liste paginée"],
        ["GET", "/api/promotions/:id", "products:read", "Détail"],
        ["POST", "/api/promotions", "products:write", "Créer"],
        ["PATCH", "/api/promotions/:id", "products:write", "Modifier"],
        ["PATCH", "/api/promotions/:id/toggle-active", "products:write", "Activer/désactiver"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Mouvements de stock — /api/stock-movements", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/stock-movements", "products:read", "Liste paginée"],
        ["GET", "/api/stock-movements/:id", "products:read", "Détail"],
        ["POST", "/api/stock-movements", "products:write", "Créer"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Boutiques — /api/shops", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/shops", "—", "Liste paginée"],
        ["GET", "/api/shops/stats", "stats:global", "Stats globales"],
        ["GET", "/api/shops/me", "shops:manage", "Ma boutique"],
        ["PATCH", "/api/shops/me", "shops:manage", "Modifier ma boutique"],
        ["POST", "/api/shops/me/logo", "shops:manage", "Uploader le logo"],
        ["POST", "/api/shops/me/images", "shops:manage", "Ajouter des images"],
        ["DELETE", "/api/shops/me/images", "shops:manage", "Supprimer une image"],
        ["GET", "/api/shops/:id", "—", "Détail d'une boutique"],
        ["PATCH", "/api/shops/:id/validate", "shops:validate", "Valider une boutique"],
        ["PATCH", "/api/shops/:id/suspend", "shops:suspend", "Suspendre une boutique"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Membres de boutique — /api/shops/me/members", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/shops/me/members", "members:read", "Liste des membres"],
        ["PATCH", "/api/shops/me/members/:id", "members:manage", "Modifier un membre"],
        ["PATCH", "/api/shops/me/members/:id/toggle-active", "members:manage", "Activer/désactiver"],
        ["DELETE", "/api/shops/me/members/:id", "members:manage", "Retirer un membre"],
        ["POST", "/api/shops/me/members/invitations", "members:manage", "Inviter un membre"],
        ["GET", "/api/shops/me/members/invitations", "members:read", "Liste des invitations"],
        ["DELETE", "/api/shops/me/members/invitations/:id", "members:manage", "Annuler une invitation"],
    ],
    [2, 8, 3, 5],
)

doc.add_heading("Commandes — /api/order", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/order/stats", "stats:read", "Statistiques commandes"],
        ["GET", "/api/order/customer", "—", "Mes commandes (client)"],
        ["POST", "/api/order/customer/checkout", "—", "Passer commande (checkout)"],
        ["GET", "/api/order/customer/:id", "—", "Détail de ma commande"],
        ["GET", "/api/order", "orders:read", "Commandes boutique"],
        ["GET", "/api/order/:id", "orders:read", "Détail commande boutique"],
        ["PATCH", "/api/order/:id/confirm", "orders:manage", "Confirmer"],
        ["PATCH", "/api/order/:id/refuse", "orders:manage", "Refuser"],
        ["PATCH", "/api/order/:id/cancel", "orders:manage", "Annuler"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Panier — /api/cart", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/cart", "cart:manage", "Voir le panier"],
        ["POST", "/api/cart/add", "cart:manage", "Ajouter un article"],
        ["POST", "/api/cart/remove", "cart:manage", "Retirer un article"],
        ["POST", "/api/cart/validate", "cart:manage", "Valider le panier"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Avis — /api/reviews", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["POST", "/api/reviews", "reviews:write", "Créer un avis"],
        ["GET", "/api/reviews/mine/:shopId", "—", "Mon avis pour une boutique"],
    ],
    [2, 7, 3, 6],
)

doc.add_heading("Rôles — /api/roles", level=3)
add_styled_table(
    ["Méthode", "URL", "Permission", "Description"],
    [
        ["GET", "/api/roles/select", "—", "Liste des rôles (select)"],
    ],
    [2, 7, 3, 6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. INFORMATIONS DU PROJET
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Informations du projet", level=1)

doc.add_heading("4.1. Lien Git et branches", level=2)

p = doc.add_paragraph()
run = p.add_run("Dépôt Git : ")
run.bold = True
run = p.add_run("https://github.com/Kevin-rm/m1p13mean-Diary-Kevin")
run.font.color.rgb = PRIMARY

doc.add_paragraph()

add_styled_table(
    ["Branche", "Description"],
    [
        ["main", "Branche principale stable"],
        ["dev", "Branche de développement"],
    ],
    [6, 12],
)

doc.add_heading("4.2. Membres du groupe", level=2)

add_styled_table(
    ["N°", "Nom et Prénom(s)", "N° Étudiant"],
    [
        ["1", "RAJOSOA Tsiory Diary Luc", "ETU002665"],
        ["2", "RAMAROZATOVO Tahiry Kevin", "ETU002546"],
    ],
    [2, 8, 4],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fin du document —")
run.font.color.rgb = GRAY
run.italic = True

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════

output_path = r"C:\Users\Kevin\Workspace\study\M1\web-avancé\Projet final - MEAN\m1p13mean-Diary-Kevin\MallHub_Documentation.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
